#!/usr/bin/env python3
"""
Convert all templates/_files/*.txt → templates/_files/*.docx
Run: python3 scripts/txt-to-docx.py
"""
import os
import re
import glob
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FILES_DIR = os.path.join(ROOT, "templates", "_files")

# Lines that look like section headings: "1. SOMETHING", "2.1 SOMETHING", "PART A — ...", "USAGE NOTES:"
HEADING_RE = re.compile(r"^(\d+\.\s+[A-Z][A-Z0-9 ,/&\-]+|PART [A-Z][^a-z]*|SCHEDULE \d+[^a-z]*|USAGE NOTES:.*|---.*|=+.*)$")
# Top-of-document title — first line, all caps, often multi-word
SUBHEADING_RE = re.compile(r"^\d+\.\d+\s+")  # 2.1, 3.4, etc.


def is_separator(line):
    s = line.strip()
    return s and (set(s) <= {"-", "="} or s.startswith("===") or s.startswith("---"))


def convert(txt_path, docx_path):
    with open(txt_path, "r", encoding="utf-8") as f:
        raw = f.read()

    lines = raw.split("\n")
    doc = Document()

    # Set base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Set narrower margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # First non-empty line is title
    title_idx = next((i for i, l in enumerate(lines) if l.strip()), 0)
    title_line = lines[title_idx].strip()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title_p.add_run(title_line)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x0A, 0x25, 0x40)

    # Process remaining lines
    i = title_idx + 1
    n = len(lines)
    in_usage = False

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Skip separators and the trailing footer/usage block start markers
        if is_separator(line):
            if "---" in stripped and not in_usage:
                in_usage = True
            i += 1
            continue

        if not stripped:
            # Preserve blank line
            doc.add_paragraph()
            i += 1
            continue

        # Detect headings — ALL-CAPS lines (>3 chars, mostly alpha/space/digits)
        is_heading = (
            len(stripped) > 3
            and stripped == stripped.upper()
            and any(c.isalpha() for c in stripped)
            and not stripped.startswith("_")
            and not stripped.startswith("(")
            and "[" not in stripped[:6]  # not a placeholder
            and len([c for c in stripped if c.isalpha()]) > 2
        )

        if is_heading and not in_usage:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x0D, 0x45, 0x67)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        else:
            # Regular paragraph — preserve leading whitespace by detecting indent
            indent_match = re.match(r"^(\s+)", line)
            indent = len(indent_match.group(1)) if indent_match else 0

            p = doc.add_paragraph()
            if indent >= 8:
                p.paragraph_format.left_indent = Inches(0.5)
            elif indent >= 4:
                p.paragraph_format.left_indent = Inches(0.25)

            # Render the (left-stripped) line, preserving internal spacing
            text = line.rstrip()
            if indent > 0:
                text = text[indent:]

            run = p.add_run(text)

            # Style usage notes block in a smaller, gray font
            if in_usage:
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        i += 1

    doc.save(docx_path)


def main():
    txt_files = sorted(glob.glob(os.path.join(FILES_DIR, "*.txt")))
    if not txt_files:
        print("No .txt files found.")
        return

    for txt_path in txt_files:
        docx_path = txt_path.replace(".txt", ".docx")
        convert(txt_path, docx_path)
        print(f"Wrote {docx_path}")

    print(f"\nDone. {len(txt_files)} files converted.")


if __name__ == "__main__":
    main()
